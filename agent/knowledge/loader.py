"""
P3: 多格式 loader。

按文件扩展名分发到具体 loader：
  .md / .markdown → TextLoader（轻量；V1 不强依赖 unstructured）
  .html / .htm    → BSHTMLLoader
  .pdf            → PyMuPDFLoader（fallback PyPDFLoader）
  .csv            → CSVLoader
  .json / .jsonl  → JSONLoader
  .docx           → _DocxCustomLoader（python-docx）
  .xlsx           → _XlsxCustomLoader（openpyxl）
  .txt / 其他     → TextLoader

业务方只配 `format: docx`（强制）或 `format: auto`（按扩展名推断）。
"""
from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document


# ────────────────────────────────────────────────────────────
# V3: 自定义 docx / xlsx loader（langchain 没现成好用的，30 行自写）
# ────────────────────────────────────────────────────────────

class _DocxCustomLoader:
    """python-docx 抽段落 + 表格，每段 / 每个表格一条 Document。"""

    def __init__(self, path: Path):
        self.path = path

    def load(self) -> list[Document]:
        from docx import Document as DocxDocument
        d = DocxDocument(str(self.path))
        out: list[Document] = []
        # 段落
        for idx, para in enumerate(d.paragraphs):
            txt = para.text.strip()
            if not txt:
                continue
            out.append(Document(
                page_content=txt,
                metadata={"file_name": self.path.name, "para_index": idx, "kind": "paragraph"},
            ))
        # 表格
        for tbl_idx, table in enumerate(d.tables):
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(" | ".join(cells))
            txt = "\n".join(rows)
            if not txt.strip():
                continue
            out.append(Document(
                page_content=f"## Table {tbl_idx + 1}\n\n{txt}",
                metadata={"file_name": self.path.name, "table_index": tbl_idx, "kind": "table"},
            ))
        return out


class _XlsxCustomLoader:
    """openpyxl 读每个 sheet，每行一条 Document。"""

    def __init__(self, path: Path):
        self.path = path

    def load(self) -> list[Document]:
        from openpyxl import load_workbook
        wb = load_workbook(str(self.path), read_only=True, data_only=True)
        out: list[Document] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers: list[str] = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                cells = [str(c) if c is not None else "" for c in row]
                if row_idx == 1:
                    headers = cells
                    continue
                line = " | ".join(
                    f"{h}: {v}" if h else v
                    for h, v in zip(headers, cells)
                )
                if line.strip():
                    out.append(Document(
                        page_content=line,
                        metadata={
                            "file_name": self.path.name,
                            "sheet": sheet_name,
                            "row": row_idx,
                        },
                    ))
        wb.close()
        return out


def _loader_for(path: Path, fmt: str):
    if fmt == "md":
        from langchain_community.document_loaders import TextLoader
        return TextLoader(str(path), encoding="utf-8")
    if fmt == "html":
        from langchain_community.document_loaders import BSHTMLLoader
        return BSHTMLLoader(str(path))
    if fmt == "pdf":
        # V2: 优先 PyMuPDFLoader（解决双栏 PDF + 页眉/页脚/页码污染），fallback PyPDFLoader
        try:
            from langchain_community.document_loaders import PyMuPDFLoader
            return PyMuPDFLoader(str(path))
        except ImportError:
            from langchain_community.document_loaders import PyPDFLoader
            return PyPDFLoader(str(path))
    if fmt == "csv":
        from langchain_community.document_loaders import CSVLoader
        return CSVLoader(str(path))
    if fmt in ("json", "jsonl"):
        from langchain_community.document_loaders import JSONLoader
        return JSONLoader(str(path), jq_schema=".", text_content=False, json_lines=(fmt == "jsonl"))
    if fmt == "docx":
        # V3: python-docx 抽段落 + 表格 → Document 列表
        return _DocxCustomLoader(path)
    if fmt == "xlsx":
        # V3: openpyxl 读 sheet + 行 → Document 列表
        return _XlsxCustomLoader(path)
    from langchain_community.document_loaders import TextLoader
    return TextLoader(str(path), encoding="utf-8")


def _auto_fmt(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".")
    if suffix in ("md", "markdown"):
        return "md"
    if suffix in ("html", "htm"):
        return "html"
    if suffix == "pdf":
        return "pdf"
    if suffix == "csv":
        return "csv"
    if suffix in ("json", "jsonl"):
        return suffix
    if suffix == "docx":
        return "docx"
    if suffix == "xlsx":
        return "xlsx"
    return "txt"


def load_file(path: Path, fmt: str = "auto") -> list[Document]:
    """加载单个文件。"""
    real_fmt = fmt if fmt != "auto" else _auto_fmt(path)
    loader = _loader_for(path, real_fmt)
    return loader.load()


def load_dir(base: Path, glob: str = "**/*", fmt: str = "auto") -> list[Document]:
    """递归加载目录下所有匹配文件。"""
    docs: list[Document] = []
    for p in sorted(base.glob(glob)):
        if not p.is_file():
            continue
        if p.suffix.lower() in (".pyc",):
            continue
        try:
            docs.extend(load_file(p, fmt=fmt))
        except Exception as e:                              # noqa: BLE001
            print(f"[skip] {p.name}: {e}")
    return docs