"""V3: docx / xlsx loader 单测。"""
from pathlib import Path

from agent.knowledge.loader import _DocxCustomLoader, _XlsxCustomLoader


def test_docx_loader_basic(tmp_path: Path):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        import pytest
        pytest.skip("python-docx not installed")

    p = tmp_path / "test.docx"
    d = DocxDocument()
    d.add_heading("标题 1", level=1)
    d.add_paragraph("这是第一段。")
    d.add_paragraph("这是第二段。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    d.save(p)

    docs = _DocxCustomLoader(p).load()
    assert any("第一段" in doc.page_content for doc in docs)
    assert any("Table 1" in doc.page_content and "A1 | B1" in doc.page_content for doc in docs)
    kinds = {doc.metadata.get("kind") for doc in docs}
    assert "paragraph" in kinds and "table" in kinds


def test_xlsx_loader_basic(tmp_path: Path):
    try:
        from openpyxl import Workbook
    except ImportError:
        import pytest
        pytest.skip("openpyxl not installed")

    p = tmp_path / "test.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age"])
    ws.append(["Alice", 30])
    ws.append(["Bob", 25])
    wb.save(p)

    docs = _XlsxCustomLoader(p).load()
    assert len(docs) == 2
    assert all(doc.metadata["sheet"] == "Sheet1" for doc in docs)
    assert any("Name: Alice" in doc.page_content for doc in docs)
    assert any("Age: 30" in doc.page_content for doc in docs)